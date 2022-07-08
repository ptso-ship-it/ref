# Download pdf reference guide, crop out header and footer area, save as a plain text file, copy into microsoft word, replace bullet points with >, add ##Top and ##Bottom to word file, make local file and master file into title case (account for also Local file and Master file cases...also watch out for double spaces, i.e. local  file), make CBCR into CbCR, remove comma inbetween Local File, and CbCR; Related party to Related-party, manuallly add in Argentina's LF and MF thresholds; Make Single year into Single-year

import docx

import pandas as pd
import re
import unicodedata

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    unicode_text = '\n'.join(fullText)
    lines_unicode = unicodedata.normalize(u'NFKD', unicode_text).encode('ascii', 'ignore').decode('utf8')
    lines_clean = lines_unicode
    lines_clean = re.sub(r'[0-9]\s*http\S+',"", lines_clean)
    lines_clean = re.sub(r'[0-9]\s*www\S+',"", lines_clean)
    lines_clean = lines_clean.replace("\t", " ")
    lines_clean = lines_clean.replace("\n", " ")
    lines_clean = re.sub(r'[^\S\n\t]+', ' ', lines_clean)
    return lines_clean

def get_all_lines(infile):
    lines = []
    for result in re.findall('##Top(.*?)##Bottom', text, re.S):
        lines.append(result)
    return lines

file = 'ey-transfer-pricing-guide-17-may-2022-txt-export-pdfxchange_edited.docx'
text = getText(file)
lines = get_all_lines(text)

'''
#Output word document test
from docx import Document
document = Document()
document.add_paragraph(lines)
document.save("fulltext_full_reference_guide.docx")
'''

countries = []
for line in lines:
    countryinfo = re.findall('Name of tax authority(.*?)capacity in the jurisdiction', line, re.S)
    countries += countryinfo

country_names = ['Albania',	'Algeria',	'Angola',	'Argentina',	'Armenia',	'Australia',	'Austria',	'Azerbaijan',	'Bahrain',	'Bangladesh',	'Belarus',	'Belgium',	'Benin',	'Bolivia',	'Bosnia and Herzegovina',	'Botswana',	'Brazil',	'Bulgaria',	'Burkina Faso',	'Cambodia',	'Cameroon',	'Canada',	'Cape Verde',	'Chad',	'Chile',	'China mainland',	'Colombia',	'Congo Brazzaville',	'Costa Rica',	'Cote dIvoire',	'Croatia',	'Cyprus',	'Czech Republic',	'Democratic Republic of Congo (DRC)',	'Denmark',	'Dominican Republic',	'Ecuador',	'Egypt',	'El Salvador',	'Estonia',	'Fiji',	'Finland',	'France',	'Gabon',	'Georgia',	'Germany',	'Ghana',	'Gibraltar',	'Greece',	'Guatemala',	'Guinea',	'Honduras',	'Hong Kong',	'Hungary',	'Iceland',	'India',	'Indonesia',	'Ireland',	'Israel',	'Italy',	'Japan',	'Jordan',	'Kazakhstan',	'Kenya',	'Kosovo',	'Kuwait',	'Latvia',	'Lebanon',	'Lithuania',	'Luxembourg',	'Madagascar',	'Malawi',	'Malaysia',	'Maldives',	'Mali',	'Malta',	'Mauritania',	'Mexico',	'Mongolia',	'Montenegro',	'Morocco',	'Mozambique',	'Namibia',	'Netherlands',	'New Zealand',	'Nicaragua',	'Nigeria',	'North Macedonia',	'Norway',	'Oman',	'Pakistan',	'Panama',	'Papua New Guinea',	'Paraguay',	'Peru',	'Philippines',	'Poland',	'Portugal',	'Puerto Rico',	'Qatar',	'Republic of Serbia',	'Romania',	'Russia',	'Rwanda',	'Saudi Arabia',	'Senegal',	'Singapore',	'Slovak Republic/Slovakia',	'Slovenia',	'South Africa',	'South Korea',	'South Sudan',	'Spain',	'Sri Lanka',	'Sweden',	'Switzerland',	'Taiwan',	'Tanzania',	'Thailand',	'Togo',	'Tunisia',	'Turkey',	'United Arab Emirates',	'Uganda',	'Ukraine',	'United Kingdom',	'United States',	'Uruguay',	'Venezuela',	'Vietnam',	'Zambia',	'Zimbabwe',]

beps_implementation_parameters = "in the local regulations\?(.*?)> Coverage"
beps_coverage_parameters = "of Master File, Local File and CbCR(.*?)> Effective or expected"
beps_template_parameters = "from OECD report template or format(.*?)> Sufficiency of BEPS"
beps_penalty_protection = "report to achieve penalty protection(.*?)c\) Is"
mcaa_parameters = "on the exchange of CbCR(.*?)4. Transfer pricing"
contemp_parameters = "submitted or prepared contemporaneously\?(.*?)> Does a local branch"
# Materiality Thresholds
tpdoc_materiality_parameters = "> Transfer pricing documentation(.*?)> Master File"
mf_materiality_parameters = "> Master File(.*?)> Local File"
lf_materiality_parameters = "> Local File(.*?)> CbCR"
cbcr_materiality_parameters = "> CbCR(.*?)> Economic analysis"
# Other requirements
language_parameters = "> Local language documentation requirement(.*?)> Safe harbor"
tp_return_parameters = "> Transfer pricing-specific returns(.*?)> Related-party disclosures along"
tp_disclosure_taxreturn_parameters = "> Related-party disclosures along with corporate income tax return(.*?)> Related-party disclosures in financial"
#Deadlines
tp_disclosure_deadline_parameters = "> Other transfer pricing disclosures and return(.*?)> Master File"
deadline_first_parameters = "a\) Filing deadline(.*?)per COV"
cit_secondary_deadline_parameters = "> Corporate income tax return(.*?)> Other transfer pricing"
mf_secondary_deadline_parameters = "> Master File(.*?)> CbCR"
cbcr_secondary_deadline_parameters = "> CbCR preparation and submission(.*?)> CbCR notification"
cbcr_notification_secondary_deadline_parameters = "> CbCR notification(.*?)b\)"
tpdoc_prep_deadline_parameters = "Transfer pricing documentation/Local File preparation deadline(.*?)c\) Transfer pricing"
tpdoc_submit_deadline_parameters = "or Local File\?(.*?)> Time period"
tpdoc_request_deadline_parameters = "> Time period or deadline for submission on tax authority request(.*?)d\)"
local_comps_parameters = "Local vs. regional comparables(.*?)> Single-year"
new_search_parameters = "> Fresh benchmarking search every year vs. rollforwards and update of the financials(.*?)> Simple"
penalty_parameters = "a\) Penalty exposure(.*?)b\)"
penalty_relief_parameters = "b\) Penalty relief(.*?)10. Statute"


def processInfo(searchCriteria, column_name, countries):
    def findInfo(searchCriteria, country_string):
        result = re.findall(rf'{searchCriteria}', country_string, re.S)
        if result:
            return result[0]
        else:
            return "No match found"

    column_values = []
    for country in countries:
        column_values.append(findInfo(searchCriteria, country))
    df[column_name] = column_values

def processSecondaryInfo(first_search_criteria, second_search_criteria, secondary_column_name, countries):
    def find_first_info(first_search_criteria, country_string):
        result = re.findall(rf'{first_search_criteria}', country_string, re.S)
        if result:
            return result[0]
        else:
            return "No match found"

    column_values = []
    for country in countries:
        column_values.append(find_first_info(first_search_criteria, country))

    def find_second_info(second_search_criteria, column_values):
        second_result = re.findall(rf'{second_search_criteria}', column_values, re.S)
        if second_result:
            return second_result[0]
        else:
            return "No match found"

    secondary_column_values = []    
    for value in column_values:
        secondary_column_values.append(find_second_info(second_search_criteria, value))
    df[secondary_column_name] = secondary_column_values


df = pd.DataFrame()
df["countries"] = country_names

processInfo(beps_implementation_parameters, "Has the jursidiction adopted BEPS in the local regulations?", countries)
processInfo(beps_coverage_parameters, "Coverage in terms of Master File, Local File, and CbCR", countries)
processInfo(beps_template_parameters, "Material differences from OECD report template", countries)
processInfo(beps_penalty_protection, "Sufficiency of BEPS format for penalty protection", countries)
processInfo(mcaa_parameters, "mcaa_status", countries)
processInfo(contemp_parameters, "Contemporaneous Doc Requirement", countries)
processInfo(tpdoc_materiality_parameters, "TP Doc Thresholds", countries)
processInfo(mf_materiality_parameters, "MF Threshold", countries)
processInfo(lf_materiality_parameters, "LF Thresholds", countries)
processInfo(cbcr_materiality_parameters, "CbCR Thresholds", countries)
processInfo(language_parameters, "Language Requirements", countries)
processInfo(tp_return_parameters, "TP Forms", countries)
processInfo(tp_disclosure_taxreturn_parameters, "TP Disclosure with Tax Return", countries)
processInfo(tp_disclosure_deadline_parameters, "TP Forms and Disclosure Deadline", countries)
processSecondaryInfo(deadline_first_parameters, cit_secondary_deadline_parameters, "CIT Deadline", countries)
processSecondaryInfo(deadline_first_parameters, mf_secondary_deadline_parameters, "MF Deadline", countries)
processSecondaryInfo(deadline_first_parameters, cbcr_secondary_deadline_parameters,"CbCR Deadlines", countries)
processSecondaryInfo(deadline_first_parameters, cbcr_notification_secondary_deadline_parameters,"CbCR Notification Deadlines", countries)
processInfo(tpdoc_prep_deadline_parameters, "TP Doc Preparation Deadline", countries)
processInfo(tpdoc_submit_deadline_parameters, "TP Doc Submit Deadline", countries)
processInfo(tpdoc_request_deadline_parameters, "Submit Upon Request Timeline", countries)
processInfo(local_comps_parameters, "Local vs Regional Comparables", countries)
processInfo(new_search_parameters, "Are New Searches Required Each Year?", countries)
processInfo(penalty_parameters, "Penalties", countries)
processInfo(penalty_relief_parameters, "Penalty Relief", countries)

df.to_csv('reference_guide_output.csv', index =False)


